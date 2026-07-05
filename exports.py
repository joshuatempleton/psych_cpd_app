from __future__ import annotations

import io
import zipfile
from typing import Any

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt

from calculations import build_insights, compute_all_cycle_summaries, compute_direct_client_contact_by_cycle, compute_metrics, compute_registrar_metrics, compute_targets
from utils import cycle_year_end_for_date, cpd_cycle_label, format_goal_links, goal_title_map


def _rows_with_cycle(rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    enriched: list[dict[str, Any]] = []
    for row in rows:
        copied = dict(row)
        year_end = cycle_year_end_for_date(copied.get("date"))
        copied["cpd_cycle_year_end"] = year_end or ""
        copied["cpd_cycle_label"] = cpd_cycle_label(year_end) if year_end else ""
        enriched.append(copied)
    return enriched


def _annual_combined_cpd_rows(portfolio: dict[str, Any]) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for row in portfolio.get("cpd_entries", []):
        copied = dict(row)
        copied["annual_cpd_source"] = "General CPD log"
        copied["annual_cpd_hours"] = copied.get("hours", 0)
        rows.append(copied)

    for row in portfolio.get("registrar_cpd_entries", []):
        copied = dict(row)
        copied["annual_cpd_source"] = "Registrar active CPD"
        copied["annual_cpd_hours"] = copied.get("hours", 0)
        copied["activity_type"] = "Registrar active CPD"
        copied["counts_to_annual_cpd"] = copied.get("counts_towards_annual_cpd", True)
        rows.append(copied)

    return _rows_with_cycle(rows)


def _annual_combined_peer_rows(portfolio: dict[str, Any]) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for row in portfolio.get("peer_entries", []):
        copied = dict(row)
        copied["annual_peer_source"] = "Peer consultation log"
        copied["annual_peer_hours"] = copied.get("own_practice_hours", 0)
        rows.append(copied)

    for row in portfolio.get("supervision_entries", []):
        copied = dict(row)
        copied["annual_peer_source"] = "Registrar supervision"
        copied["annual_peer_hours"] = copied.get("hours", 0)
        copied["counts_to_annual_peer_consultation"] = copied.get("counts_towards_peer_consultation", True)
        rows.append(copied)

    return _rows_with_cycle(rows)


def export_csv_zip(portfolio: dict[str, Any]) -> bytes:
    out = io.BytesIO()

    tables = {
        "learning_goals.csv": portfolio.get("learning_goals", []),
        "general_cpd_entries_all_years.csv": _rows_with_cycle(portfolio.get("cpd_entries", [])),
        "registrar_active_cpd_entries_program_total.csv": _rows_with_cycle(portfolio.get("registrar_cpd_entries", [])),
        "annual_cpd_combined_documentation.csv": _annual_combined_cpd_rows(portfolio),
        "peer_consultation_entries_all_years.csv": _rows_with_cycle(portfolio.get("peer_entries", [])),
        "registrar_supervision_entries_program_total.csv": _rows_with_cycle(portfolio.get("supervision_entries", [])),
        "registrar_practice_log_entries_program_total.csv": _rows_with_cycle(portfolio.get("registrar_practice_entries", [])),
        "registrar_direct_client_contact_by_cycle.csv": compute_direct_client_contact_by_cycle(portfolio),
        "annual_peer_consultation_combined_documentation.csv": _annual_combined_peer_rows(portfolio),
        "annual_cpd_cycle_summaries.csv": compute_all_cycle_summaries(portfolio),
        "registrar_competencies.csv": portfolio.get("competency_assessments", []),
        "registrar_deadlines.csv": portfolio.get("registrar_deadlines", []),
    }

    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zf:
        for filename, rows in tables.items():
            zf.writestr(filename, pd.DataFrame(rows).to_csv(index=False))

        metrics = compute_metrics(portfolio)
        reg = compute_registrar_metrics(portfolio)
        zf.writestr(
            "selected_cycle_and_registrar_summary.csv",
            pd.DataFrame(
                [
                    {"Metric": "Selected annual CPD cycle", "Value": metrics["cycle_label"]},
                    {"Metric": "Selected cycle total CPD hours", "Value": metrics["total_hours"]},
                    {"Metric": "Selected cycle general CPD hours - total", "Value": metrics["cpd_hours"]},
                    {"Metric": "Selected cycle general CPD from general log", "Value": metrics["general_cpd_hours"]},
                    {"Metric": "Selected cycle general CPD from registrar active CPD", "Value": metrics["registrar_active_cpd_hours_in_cycle"]},
                    {"Metric": "Selected cycle peer consultation hours - total", "Value": metrics["peer_hours"]},
                    {"Metric": "Selected cycle peer hours from peer log", "Value": metrics["standalone_peer_hours"]},
                    {"Metric": "Selected cycle peer hours from registrar supervision", "Value": metrics["registrar_supervision_peer_hours_in_cycle"]},
                    {"Metric": "Registrar practice hours from practice log - program total", "Value": reg["practice_hours"]},
                    {"Metric": "Registrar practice log entries", "Value": reg["practice_log_entry_count"]},
                    {"Metric": "Direct client contact hours - selected cycle", "Value": reg["direct_client_contact_hours_selected_cycle"]},
                    {"Metric": "Direct client contact minimum - selected cycle", "Value": 176.0},
                    {"Metric": "Registrar supervision hours - program total", "Value": reg["supervision_hours"]},
                    {"Metric": "Registrar active CPD hours - program total", "Value": reg["active_cpd_hours"]},
                ]
            ).to_csv(index=False),
        )

    out.seek(0)
    return out.getvalue()


def _add_table(doc: Document, headers: list[str], rows: list[list[Any]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row_data in rows:
        row = table.add_row().cells
        for i, value in enumerate(row_data):
            row[i].text = "" if value is None else str(value)


def export_docx(portfolio: dict[str, Any]) -> bytes:
    metrics = compute_metrics(portfolio)
    goals_map = goal_title_map(portfolio["learning_goals"])

    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(0.6)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    title = doc.add_heading("Psychology CPD Portfolio Export", level=0)
    title.alignment = 1

    profile = portfolio["profile"]
    targets = compute_targets(profile)

    doc.add_heading("Portfolio details", level=1)
    _add_table(
        doc,
        ["Field", "Value"],
        [
            ["Psychologist name", profile.get("psychologist_name", "")],
            ["Selected annual CPD cycle", targets["cycle_label"]],
            ["Endorsements", ", ".join(profile.get("endorsements", []))],
            ["Registrar", "Yes" if profile.get("is_registrar") else "No"],
            ["Registrar area", portfolio.get("registrar", {}).get("area", "")],
        ],
    )

    doc.add_heading("Summary insights", level=1)
    for line in (portfolio.get("summary_insights") or build_insights(portfolio, metrics)).splitlines():
        doc.add_paragraph(line)

    doc.add_heading("Learning plan", level=1)
    _add_table(
        doc,
        ["Goal", "Learning need", "Proposed activities", "Dates", "Anticipated outcomes", "Review date", "Outcomes achieved"],
        [
            [
                g.get("title"),
                g.get("learning_need"),
                g.get("proposed_activities"),
                g.get("proposed_dates"),
                g.get("anticipated_outcomes"),
                g.get("review_date"),
                g.get("outcomes_achieved"),
            ]
            for g in portfolio.get("learning_goals", [])
        ],
    )

    doc.add_heading("Annual CPD documentation - combined", level=1)
    _add_table(
        doc,
        ["Date", "Source", "Type", "Details", "Endorsement area", "Hours", "Evidence", "Reflection"],
        [
            [
                e.get("date"),
                e.get("annual_cpd_source"),
                e.get("activity_type"),
                e.get("activity_details"),
                e.get("endorsement_area"),
                e.get("annual_cpd_hours"),
                f"{e.get('evidence_type', '')}: {e.get('evidence_details', '')}",
                e.get("reflection"),
            ]
            for e in _annual_combined_cpd_rows(portfolio)
            if e.get("counts_to_annual_cpd", True)
        ],
    )

    doc.add_heading("Annual peer consultation documentation - combined", level=1)
    _add_table(
        doc,
        ["Date", "Source", "Format / Type", "Focus / Notes", "People involved", "Hours", "Evidence / where stored"],
        [
            [
                e.get("date"),
                e.get("annual_peer_source"),
                e.get("format") or e.get("supervision_type"),
                e.get("focus") or e.get("notes"),
                e.get("colleagues") or e.get("supervisor_name"),
                e.get("annual_peer_hours"),
                e.get("evidence") or f"{e.get('evidence_type', '')}: {e.get('evidence_details', '')}",
            ]
            for e in _annual_combined_peer_rows(portfolio)
            if e.get("counts_to_annual_peer_consultation", True)
        ],
    )

    doc.add_heading("General CPD log only", level=1)
    _add_table(
        doc,
        ["Date", "Type", "Details", "Area", "Hours", "Goals", "Evidence", "Reflection"],
        [
            [
                e.get("date"),
                e.get("activity_type"),
                e.get("activity_details"),
                e.get("area_of_practice"),
                e.get("hours"),
                format_goal_links(e.get("goal_ids", []), goals_map),
                f"{e.get('evidence_type', '')}: {e.get('evidence_details', '')}",
                e.get("reflection"),
            ]
            for e in portfolio.get("cpd_entries", [])
        ],
    )

    doc.add_heading("Registrar active CPD log", level=1)
    _add_table(
        doc,
        ["Date", "Endorsement area", "Activity", "Hours", "Competency domains", "Counts to annual CPD", "Evidence", "Reflection"],
        [
            [
                e.get("date"),
                e.get("endorsement_area"),
                e.get("activity_details"),
                e.get("hours"),
                "; ".join(e.get("competency_domains", [])),
                "Yes" if e.get("counts_towards_annual_cpd", True) else "No",
                f"{e.get('evidence_type', '')}: {e.get('evidence_details', '')}",
                e.get("reflection"),
            ]
            for e in portfolio.get("registrar_cpd_entries", [])
        ],
    )

    doc.add_heading("Peer consultation log only", level=1)
    _add_table(
        doc,
        ["Date", "Format", "Focus", "Colleagues", "Total hours", "Own-practice hours", "Area", "Goals", "Evidence", "Reflection"],
        [
            [
                e.get("date"),
                e.get("format"),
                e.get("focus"),
                e.get("colleagues"),
                e.get("total_hours"),
                e.get("own_practice_hours"),
                e.get("area_of_practice"),
                format_goal_links(e.get("goal_ids", []), goals_map),
                f"{e.get('evidence_type', '')}: {e.get('evidence_details', '')}",
                e.get("reflection"),
            ]
            for e in portfolio.get("peer_entries", [])
        ],
    )

    doc.add_heading("Registrar practice log", level=1)
    _add_table(
        doc,
        ["Date", "Practice hours", "Direct client contact hours"],
        [
            [
                e.get("date"),
                e.get("practice_hours"),
                e.get("direct_client_contact_hours"),
            ]
            for e in portfolio.get("registrar_practice_entries", [])
        ],
    )

    doc.add_heading("Registrar direct client contact by CPD cycle", level=1)
    _add_table(
        doc,
        ["CPD cycle", "Direct client contact hours", "Minimum required", "Remaining", "Requirement met"],
        [
            [
                e.get("cycle_label"),
                e.get("direct_client_contact_hours"),
                e.get("minimum_required_hours"),
                e.get("remaining_hours"),
                "Yes" if e.get("requirement_met") else "No",
            ]
            for e in compute_direct_client_contact_by_cycle(portfolio)
        ],
    )

    doc.add_heading("Registrar supervision log", level=1)
    _add_table(
        doc,
        ["Date", "Hours", "Practice hours", "Supervisor", "Type", "Competency domains", "Counts to annual peer consultation", "Notes", "Evidence"],
        [
            [
                e.get("date"),
                e.get("hours"),
                e.get("practice_hours"),
                e.get("supervisor_name"),
                e.get("supervision_type"),
                "; ".join(e.get("competency_domains", [])),
                "Yes" if e.get("counts_towards_peer_consultation", True) else "No",
                e.get("notes"),
                e.get("evidence"),
            ]
            for e in portfolio.get("supervision_entries", [])
        ],
    )

    doc.add_heading("Endorsement competencies", level=1)
    _add_table(
        doc,
        ["Endorsement area", "Domain", "Subdomain", "Rating", "Last reviewed", "Evidence / comments"],
        [
            [
                e.get("endorsement_area"),
                e.get("domain"),
                e.get("subdomain"),
                e.get("rating"),
                e.get("last_reviewed"),
                e.get("evidence"),
            ]
            for e in portfolio.get("competency_assessments", [])
        ],
    )

    doc.add_heading("Registrar deadlines", level=1)
    _add_table(
        doc,
        ["Type", "Due date", "Status", "Notes"],
        [[e.get("type"), e.get("due_date"), e.get("status"), e.get("notes")] for e in portfolio.get("registrar_deadlines", [])],
    )

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()
