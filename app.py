import io
import json
import uuid
from copy import deepcopy
from datetime import date, datetime
from typing import Any, Dict, List, Optional

import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


APP_TITLE = "Psychologist CPD Portfolio Tracker"
APP_VERSION = "1.1.0"

CPD_CYCLE_START_MONTH = 12
CPD_CYCLE_START_DAY = 1
CPD_CYCLE_END_MONTH = 11
CPD_CYCLE_END_DAY = 30

EVIDENCE_OPTIONS = [
    "Certificate of attendance/completion",
    "Receipt",
    "Reading list",
    "Professional association CPD record",
    "Degree certificate / academic transcript",
    "Assignment / thesis / research report / published article",
    "Supervision / mentoring plan or progress report",
    "Flyer / program / registration confirmation",
    "Self-recorded reflection only",
    "Other",
]

CPD_ACTIVITY_TYPES = [
    "Workshop",
    "Seminar",
    "Conference",
    "Reading",
    "Online course",
    "Webinar",
    "Teaching / presenting",
    "Research / writing",
    "Supervision-related training",
    "Other",
]

PEER_FORMATS = [
    "Individual supervision",
    "Mentoring",
    "Consultation",
    "Group peer consultation",
    "Case discussion",
    "Peer support network",
    "Learning plan review",
    "Other",
]

ENDORSEMENT_OPTIONS = [
    "Clinical psychology",
    "Counselling psychology",
    "Educational and developmental psychology",
    "Forensic psychology",
    "Health psychology",
    "Organisational psychology",
    "Sport and exercise psychology",
    "Community psychology",
    "Other / custom",
]


def today_str() -> str:
    return date.today().isoformat()


def new_id() -> str:
    return uuid.uuid4().hex[:10]


def cpd_cycle_label(year_end: int) -> str:
    start_year = year_end - 1
    return f"1 Dec {start_year} to 30 Nov {year_end}"


def in_cpd_cycle(entry_date: Optional[str], cycle_year_end: int) -> bool:
    if not entry_date:
        return False
    try:
        d = datetime.strptime(entry_date, "%Y-%m-%d").date()
    except ValueError:
        return False
    start = date(cycle_year_end - 1, 12, 1)
    end = date(cycle_year_end, 11, 30)
    return start <= d <= end


def safe_float(v: Any) -> float:
    try:
        return round(float(v), 2)
    except Exception:
        return 0.0


def parse_iso_date(value: Any) -> date:
    if isinstance(value, date):
        return value
    if isinstance(value, str):
        try:
            return datetime.strptime(value, "%Y-%m-%d").date()
        except ValueError:
            return date.today()
    return date.today()


def default_profile() -> Dict[str, Any]:
    current_year = date.today().year
    year_end = current_year if date.today() <= date(current_year, 11, 30) else current_year + 1
    return {
        "psychologist_name": "",
        "registration_cycle_year_end": year_end,
        "full_year_general_registration": True,
        "months_general_registration": 12,
        "has_board_exemption": False,
        "exemption_notes": "",
        "is_part_time": False,
        "has_endorsement": False,
        "endorsements": [],
        "is_registrar": False,
        "registrar_area": "",
        "is_board_approved_supervisor": False,
        "supervisor_last_training_date": "",
        "practice_context": "",
        "position_description_note": "",
        "signature_name": "",
    }


def default_portfolio() -> Dict[str, Any]:
    return {
        "meta": {
            "app_title": APP_TITLE,
            "app_version": APP_VERSION,
            "created_at": datetime.now().isoformat(),
            "last_saved_at": None,
            "portfolio_id": new_id(),
        },
        "profile": default_profile(),
        "learning_goals": [
            {
                "id": new_id(),
                "title": "Goal 1",
                "learning_need": "",
                "proposed_activities": "",
                "proposed_dates": "",
                "anticipated_outcomes": "",
                "review_date": "",
                "outcomes_achieved": "",
                "status": "Planned",
            },
            {
                "id": new_id(),
                "title": "Goal 2",
                "learning_need": "",
                "proposed_activities": "",
                "proposed_dates": "",
                "anticipated_outcomes": "",
                "review_date": "",
                "outcomes_achieved": "",
                "status": "Planned",
            },
            {
                "id": new_id(),
                "title": "Goal 3",
                "learning_need": "",
                "proposed_activities": "",
                "proposed_dates": "",
                "anticipated_outcomes": "",
                "review_date": "",
                "outcomes_achieved": "",
                "status": "Planned",
            },
        ],
        "cpd_entries": [],
        "peer_entries": [],
        "summary_insights": "",
    }


def ensure_state() -> None:
    if "portfolio" not in st.session_state:
        st.session_state.portfolio = default_portfolio()
    if "editing_cpd_id" not in st.session_state:
        st.session_state.editing_cpd_id = None
    if "editing_peer_id" not in st.session_state:
        st.session_state.editing_peer_id = None


def goal_title_map(goals: List[Dict[str, Any]]) -> Dict[str, str]:
    return {g["id"]: (g.get("title") or "Untitled goal") for g in goals}


def get_entry_by_id(entries: List[Dict[str, Any]], entry_id: Optional[str]) -> Optional[Dict[str, Any]]:
    if not entry_id:
        return None
    for entry in entries:
        if entry.get("id") == entry_id:
            return entry
    return None


def upsert_entry(entries: List[Dict[str, Any]], new_entry: Dict[str, Any]) -> None:
    for idx, entry in enumerate(entries):
        if entry.get("id") == new_entry.get("id"):
            entries[idx] = new_entry
            return
    entries.append(new_entry)


def delete_entry(entries: List[Dict[str, Any]], entry_id: str) -> None:
    entries[:] = [e for e in entries if e.get("id") != entry_id]


def compute_targets(profile: Dict[str, Any]) -> Dict[str, Any]:
    full_year = bool(profile.get("full_year_general_registration", True))
    months = int(profile.get("months_general_registration", 12) or 12)
    months = max(0, min(months, 12))

    if full_year:
        total_target = 30.0
        peer_target = 10.0
        general_target = 20.0
    else:
        total_target = round(months * 2.5, 2)
        peer_target = round(months * (50 / 60), 2)
        general_target = round(months * (100 / 60), 2)

    endorsement_rules = {}
    endorsements = [e for e in profile.get("endorsements", []) if e]
    count = len(endorsements)
    if count == 1:
        endorsement_rules[endorsements[0]] = 16.0
    elif count == 2:
        endorsement_rules[endorsements[0]] = 15.0
        endorsement_rules[endorsements[1]] = 15.0
    elif count >= 3:
        split = round(30.0 / count, 2)
        for e in endorsements:
            endorsement_rules[e] = split

    return {
        "total_target": total_target,
        "peer_target": peer_target,
        "general_target": general_target,
        "endorsement_targets": endorsement_rules,
        "cycle_label": cpd_cycle_label(int(profile["registration_cycle_year_end"])),
    }


def compute_metrics(portfolio: Dict[str, Any]) -> Dict[str, Any]:
    profile = portfolio["profile"]
    targets = compute_targets(profile)
    cycle_year_end = int(profile["registration_cycle_year_end"])

    cpd_entries = portfolio["cpd_entries"]
    peer_entries = portfolio["peer_entries"]

    cpd_in_cycle = [x for x in cpd_entries if in_cpd_cycle(x.get("date"), cycle_year_end)]
    peer_in_cycle = [x for x in peer_entries if in_cpd_cycle(x.get("date"), cycle_year_end)]

    cpd_hours = round(sum(safe_float(x.get("hours")) for x in cpd_in_cycle), 2)
    peer_hours = round(sum(safe_float(x.get("own_practice_hours")) for x in peer_in_cycle), 2)
    total_hours = round(cpd_hours + peer_hours, 2)

    endorsement_hours: Dict[str, float] = {}
    for entry in cpd_in_cycle:
        area = entry.get("area_of_practice") or "Unspecified"
        endorsement_hours[area] = round(endorsement_hours.get(area, 0.0) + safe_float(entry.get("hours")), 2)
    for entry in peer_in_cycle:
        area = entry.get("area_of_practice") or "Unspecified"
        endorsement_hours[area] = round(
            endorsement_hours.get(area, 0.0) + safe_float(entry.get("own_practice_hours")),
            2,
        )

    return {
        **targets,
        "cpd_hours": cpd_hours,
        "peer_hours": peer_hours,
        "total_hours": total_hours,
        "cpd_remaining": round(max(0.0, targets["general_target"] - cpd_hours), 2),
        "peer_remaining": round(max(0.0, targets["peer_target"] - peer_hours), 2),
        "total_remaining": round(max(0.0, targets["total_target"] - total_hours), 2),
        "cpd_in_cycle": cpd_in_cycle,
        "peer_in_cycle": peer_in_cycle,
        "endorsement_hours": endorsement_hours,
    }


def portfolio_json_bytes(portfolio: Dict[str, Any]) -> bytes:
    payload = deepcopy(portfolio)
    payload["meta"]["last_saved_at"] = datetime.now().isoformat()
    return json.dumps(payload, indent=2, ensure_ascii=False).encode("utf-8")


def load_portfolio_from_bytes(raw: bytes) -> Dict[str, Any]:
    data = json.loads(raw.decode("utf-8"))
    if not isinstance(data, dict):
        raise ValueError("Uploaded file is not a valid portfolio JSON object.")
    required = {"meta", "profile", "learning_goals", "cpd_entries", "peer_entries", "summary_insights"}
    missing = required - set(data.keys())
    if missing:
        raise ValueError(f"Uploaded file is missing required keys: {', '.join(sorted(missing))}")
    return data


def format_goal_links(goal_ids: List[str], goals_map: Dict[str, str]) -> str:
    if not goal_ids:
        return ""
    return "; ".join(goals_map.get(gid, gid) for gid in goal_ids)


def build_insights(portfolio: Dict[str, Any], metrics: Dict[str, Any]) -> str:
    profile = portfolio["profile"]
    targets = compute_targets(profile)

    lines = [
        f"Registration cycle: {targets['cycle_label']}",
        f"Total logged hours in cycle: {metrics['total_hours']} / {targets['total_target']}",
        f"General CPD logged: {metrics['cpd_hours']} / {targets['general_target']}",
        f"Peer consultation logged: {metrics['peer_hours']} / {targets['peer_target']}",
    ]

    if metrics["total_remaining"] > 0:
        lines.append(f"Remaining total hours: {metrics['total_remaining']}")
    else:
        lines.append("Total CPD target has been met or exceeded.")

    if metrics["peer_remaining"] > 0:
        lines.append(f"Remaining peer consultation hours: {metrics['peer_remaining']}")
    else:
        lines.append("Peer consultation target has been met or exceeded.")

    endorsements = targets["endorsement_targets"]
    if endorsements:
        lines.append("Endorsement tracking:")
        for area, target in endorsements.items():
            actual = round(metrics["endorsement_hours"].get(area, 0.0), 2)
            lines.append(f"- {area}: {actual} / {target} hours")

    if profile.get("is_board_approved_supervisor"):
        lines.append(
            "Board-approved supervisor flag is on: remember supervisor refresher/training is required at least every five years."
        )

    if profile.get("is_registrar"):
        lines.append(
            "Registrar flag is on: the dashboard only checks the base Board minimums. Registrar-specific supervision/program requirements should still be checked separately."
        )

    return "\n".join(lines)


def export_docx(portfolio: Dict[str, Any]) -> bytes:
    metrics = compute_metrics(portfolio)
    goals_map = goal_title_map(portfolio["learning_goals"])

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10)

    title = doc.add_heading("Psychology CPD Portfolio Export", level=0)
    title.alignment = 1
    p = doc.add_paragraph()
    p.add_run(
        "This export is designed to align with the Appendix E CPD forms structure and audit-style record keeping."
    ).italic = True

    profile = portfolio["profile"]
    targets = compute_targets(profile)

    doc.add_heading("Portfolio details", level=1)
    details = [
        ("Psychologist name", profile.get("psychologist_name", "")),
        ("Signature name", profile.get("signature_name", "")),
        ("Registration cycle", targets["cycle_label"]),
        ("Full-year general registration", "Yes" if profile.get("full_year_general_registration") else "No"),
        ("Months of general registration", str(profile.get("months_general_registration", ""))),
        ("Board exemption/variation noted", "Yes" if profile.get("has_board_exemption") else "No"),
        ("Exemption notes", profile.get("exemption_notes", "")),
        ("Endorsements", ", ".join(profile.get("endorsements", []))),
        ("Registrar", "Yes" if profile.get("is_registrar") else "No"),
        ("Registrar area", profile.get("registrar_area", "")),
        ("Board-approved supervisor", "Yes" if profile.get("is_board_approved_supervisor") else "No"),
        ("Supervisor training date", profile.get("supervisor_last_training_date", "")),
        ("Practice context", profile.get("practice_context", "")),
        ("Position/practice relevance note", profile.get("position_description_note", "")),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "Field"
    hdr[1].text = "Value"
    for key, value in details:
        row = table.add_row().cells
        row[0].text = str(key)
        row[1].text = str(value)

    doc.add_heading("Summary insights", level=1)
    summary_text = portfolio.get("summary_insights") or build_insights(portfolio, metrics)
    for line in summary_text.splitlines():
        doc.add_paragraph(line)

    doc.add_heading("Continuing professional development learning plan", level=1)
    lp = doc.add_table(rows=1, cols=7)
    lp.style = "Table Grid"
    headers = lp.rows[0].cells
    headers[0].text = "Goal"
    headers[1].text = "Learning needs identified and goals set"
    headers[2].text = "Type of activities proposed"
    headers[3].text = "Dates proposed activities planned"
    headers[4].text = "Outcomes anticipated"
    headers[5].text = "Review date"
    headers[6].text = "Outcomes achieved"
    for g in portfolio["learning_goals"]:
        row = lp.add_row().cells
        row[0].text = g.get("title", "")
        row[1].text = g.get("learning_need", "")
        row[2].text = g.get("proposed_activities", "")
        row[3].text = g.get("proposed_dates", "")
        row[4].text = g.get("anticipated_outcomes", "")
        row[5].text = g.get("review_date", "")
        row[6].text = g.get("outcomes_achieved", "")

    doc.add_page_break()
    doc.add_heading("Continuing professional development activity log", level=1)
    cpd = doc.add_table(rows=1, cols=8)
    cpd.style = "Table Grid"
    h = cpd.rows[0].cells
    h[0].text = "Date"
    h[1].text = "Type of activity"
    h[2].text = "Activity details"
    h[3].text = "Area of practice"
    h[4].text = "Hours claimed"
    h[5].text = "Related learning goals"
    h[6].text = "Evidence"
    h[7].text = "Reflection"
    for entry in portfolio["cpd_entries"]:
        row = cpd.add_row().cells
        evidence = entry.get("evidence_type", "")
        if entry.get("evidence_details"):
            evidence = f"{evidence}: {entry['evidence_details']}"
        row[0].text = entry.get("date", "")
        row[1].text = entry.get("activity_type", "")
        row[2].text = entry.get("activity_details", "")
        row[3].text = entry.get("area_of_practice", "")
        row[4].text = str(entry.get("hours", ""))
        row[5].text = format_goal_links(entry.get("goal_ids", []), goals_map)
        row[6].text = evidence
        row[7].text = entry.get("reflection", "")

    doc.add_heading("Peer consultation log", level=1)
    peer = doc.add_table(rows=1, cols=9)
    peer.style = "Table Grid"
    h = peer.rows[0].cells
    h[0].text = "Date"
    h[1].text = "Format"
    h[2].text = "Focus of peer consultation"
    h[3].text = "Colleagues involved"
    h[4].text = "Total duration (hours)"
    h[5].text = "Own practice hours"
    h[6].text = "Area of practice"
    h[7].text = "Related learning goals"
    h[8].text = "Evidence"
    for entry in portfolio["peer_entries"]:
        row = peer.add_row().cells
        evidence = entry.get("evidence_type", "")
        if entry.get("evidence_details"):
            evidence = f"{evidence}: {entry['evidence_details']}"
        row[0].text = entry.get("date", "")
        row[1].text = entry.get("format", "")
        row[2].text = entry.get("focus", "")
        row[3].text = entry.get("colleagues", "")
        row[4].text = str(entry.get("total_hours", ""))
        row[5].text = str(entry.get("own_practice_hours", ""))
        row[6].text = entry.get("area_of_practice", "")
        row[7].text = format_goal_links(entry.get("goal_ids", []), goals_map)
        row[8].text = evidence

    doc.add_heading("Professional development journal", level=1)
    for entry in portfolio["cpd_entries"]:
        doc.add_paragraph(f"{entry.get('date', '')} — {entry.get('activity_type', '')}", style="Heading 3")
        doc.add_paragraph(f"Activity details: {entry.get('activity_details', '')}")
        doc.add_paragraph(f"Related learning goals: {format_goal_links(entry.get('goal_ids', []), goals_map)}")
        doc.add_paragraph(f"Reflection: {entry.get('reflection', '')}")

    doc.add_heading("Peer consultation journal", level=1)
    for entry in portfolio["peer_entries"]:
        doc.add_paragraph(f"{entry.get('date', '')} — {entry.get('format', '')}", style="Heading 3")
        doc.add_paragraph(f"Focus: {entry.get('focus', '')}")
        doc.add_paragraph(f"Related learning goals: {format_goal_links(entry.get('goal_ids', []), goals_map)}")
        doc.add_paragraph(f"Reflection: {entry.get('reflection', '')}")

    doc.add_heading("Evidence checklist", level=1)
    doc.add_paragraph(
        "Store your external evidence separately and keep it with this export and your JSON portfolio file."
    )
    ev = doc.add_table(rows=1, cols=4)
    ev.style = "Table Grid"
    h = ev.rows[0].cells
    h[0].text = "Entry type"
    h[1].text = "Date"
    h[2].text = "Evidence type"
    h[3].text = "Evidence detail / where stored"
    for entry in portfolio["cpd_entries"]:
        row = ev.add_row().cells
        row[0].text = "General CPD"
        row[1].text = entry.get("date", "")
        row[2].text = entry.get("evidence_type", "")
        row[3].text = entry.get("evidence_details", "")
    for entry in portfolio["peer_entries"]:
        row = ev.add_row().cells
        row[0].text = "Peer consultation"
        row[1].text = entry.get("date", "")
        row[2].text = entry.get("evidence_type", "")
        row[3].text = entry.get("evidence_details", "")

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()


def _pdf_table(data: List[List[str]], col_widths: Optional[List[float]] = None) -> Table:
    t = Table(data, colWidths=col_widths, repeatRows=1)
    t.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D9E8FB")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("LEADING", (0, 0), (-1, -1), 10),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.lightgrey]),
            ]
        )
    )
    return t


def export_pdf(portfolio: Dict[str, Any]) -> bytes:
    metrics = compute_metrics(portfolio)
    goals_map = goal_title_map(portfolio["learning_goals"])
    out = io.BytesIO()

    doc = SimpleDocTemplate(
        out,
        pagesize=A4,
        topMargin=1.2 * cm,
        bottomMargin=1.2 * cm,
        leftMargin=1.2 * cm,
        rightMargin=1.2 * cm,
    )
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="Small", fontSize=8, leading=10))
    story = []

    story.append(Paragraph("<b>Psychology CPD Portfolio Export</b>", styles["Title"]))
    story.append(Paragraph("Structured for documentation and audit-style record keeping.", styles["Italic"]))
    story.append(Spacer(1, 0.3 * cm))

    profile = portfolio["profile"]
    targets = compute_targets(profile)

    details = [
        ["Field", "Value"],
        ["Psychologist name", profile.get("psychologist_name", "")],
        ["Registration cycle", targets["cycle_label"]],
        ["Full-year general registration", "Yes" if profile.get("full_year_general_registration") else "No"],
        ["Months of general registration", str(profile.get("months_general_registration", ""))],
        ["Endorsements", ", ".join(profile.get("endorsements", []))],
        ["Registrar", "Yes" if profile.get("is_registrar") else "No"],
        ["Board-approved supervisor", "Yes" if profile.get("is_board_approved_supervisor") else "No"],
        ["Practice context", profile.get("practice_context", "")],
    ]
    story.append(Paragraph("<b>Portfolio details</b>", styles["Heading1"]))
    story.append(_pdf_table(details, [5 * cm, 11 * cm]))
    story.append(Spacer(1, 0.3 * cm))

    summary_text = portfolio.get("summary_insights") or build_insights(portfolio, metrics)
    story.append(Paragraph("<b>Summary insights</b>", styles["Heading1"]))
    for line in summary_text.splitlines():
        story.append(Paragraph(line.replace("\n", " "), styles["BodyText"]))
    story.append(Spacer(1, 0.3 * cm))

    story.append(Paragraph("<b>Learning plan</b>", styles["Heading1"]))
    learning_table = [[
        "Goal",
        "Learning need",
        "Proposed activities",
        "Proposed dates",
        "Anticipated outcomes",
        "Review date",
        "Outcomes achieved",
    ]]
    for g in portfolio["learning_goals"]:
        learning_table.append([
            g.get("title", ""),
            g.get("learning_need", ""),
            g.get("proposed_activities", ""),
            g.get("proposed_dates", ""),
            g.get("anticipated_outcomes", ""),
            g.get("review_date", ""),
            g.get("outcomes_achieved", ""),
        ])
    story.append(_pdf_table(
        learning_table,
        [2.0 * cm, 3.2 * cm, 3.0 * cm, 2.0 * cm, 3.0 * cm, 1.8 * cm, 3.0 * cm],
    ))
    story.append(Spacer(1, 0.3 * cm))

    story.append(Paragraph("<b>CPD activity log</b>", styles["Heading1"]))
    cpd_table = [["Date", "Type", "Activity details", "Area", "Hours", "Goals", "Evidence", "Reflection"]]
    for entry in portfolio["cpd_entries"]:
        evidence = entry.get("evidence_type", "")
        if entry.get("evidence_details"):
            evidence = f"{evidence}: {entry['evidence_details']}"
        cpd_table.append([
            entry.get("date", ""),
            entry.get("activity_type", ""),
            entry.get("activity_details", ""),
            entry.get("area_of_practice", ""),
            str(entry.get("hours", "")),
            format_goal_links(entry.get("goal_ids", []), goals_map),
            evidence,
            entry.get("reflection", ""),
        ])
    story.append(_pdf_table(
        cpd_table,
        [1.5 * cm, 1.8 * cm, 3.5 * cm, 2.0 * cm, 1.2 * cm, 2.6 * cm, 2.8 * cm, 3.6 * cm],
    ))
    story.append(Spacer(1, 0.3 * cm))

    story.append(Paragraph("<b>Peer consultation log</b>", styles["Heading1"]))
    peer_table = [["Date", "Format", "Focus", "Colleagues", "Total hrs", "Own-practice hrs", "Area", "Goals", "Evidence"]]
    for entry in portfolio["peer_entries"]:
        evidence = entry.get("evidence_type", "")
        if entry.get("evidence_details"):
            evidence = f"{evidence}: {entry['evidence_details']}"
        peer_table.append([
            entry.get("date", ""),
            entry.get("format", ""),
            entry.get("focus", ""),
            entry.get("colleagues", ""),
            str(entry.get("total_hours", "")),
            str(entry.get("own_practice_hours", "")),
            entry.get("area_of_practice", ""),
            format_goal_links(entry.get("goal_ids", []), goals_map),
            evidence,
        ])
    story.append(_pdf_table(
        peer_table,
        [1.3 * cm, 1.7 * cm, 3.0 * cm, 2.2 * cm, 1.3 * cm, 1.8 * cm, 1.8 * cm, 2.4 * cm, 2.8 * cm],
    ))
    story.append(Spacer(1, 0.3 * cm))

    story.append(Paragraph("<b>Professional development journal</b>", styles["Heading1"]))
    for entry in portfolio["cpd_entries"]:
        story.append(Paragraph(f"<b>{entry.get('date', '')} — {entry.get('activity_type', '')}</b>", styles["BodyText"]))
        story.append(Paragraph(f"Activity details: {entry.get('activity_details', '')}", styles["Small"]))
        story.append(Paragraph(f"Goals: {format_goal_links(entry.get('goal_ids', []), goals_map)}", styles["Small"]))
        story.append(Paragraph(f"Reflection: {entry.get('reflection', '')}", styles["Small"]))
        story.append(Spacer(1, 0.15 * cm))

    story.append(Paragraph("<b>Peer consultation journal</b>", styles["Heading1"]))
    for entry in portfolio["peer_entries"]:
        story.append(Paragraph(f"<b>{entry.get('date', '')} — {entry.get('format', '')}</b>", styles["BodyText"]))
        story.append(Paragraph(f"Focus: {entry.get('focus', '')}", styles["Small"]))
        story.append(Paragraph(f"Goals: {format_goal_links(entry.get('goal_ids', []), goals_map)}", styles["Small"]))
        story.append(Paragraph(f"Reflection: {entry.get('reflection', '')}", styles["Small"]))
        story.append(Spacer(1, 0.15 * cm))

    story.append(Paragraph("<b>Evidence checklist</b>", styles["Heading1"]))
    story.append(Paragraph(
        "Keep external evidence stored separately with your JSON portfolio file and this export.",
        styles["Small"],
    ))
    ev_table = [["Entry type", "Date", "Evidence type", "Evidence detail / where stored"]]
    for entry in portfolio["cpd_entries"]:
        ev_table.append(["General CPD", entry.get("date", ""), entry.get("evidence_type", ""), entry.get("evidence_details", "")])
    for entry in portfolio["peer_entries"]:
        ev_table.append(["Peer consultation", entry.get("date", ""), entry.get("evidence_type", ""), entry.get("evidence_details", "")])
    story.append(_pdf_table(ev_table, [2.4 * cm, 2.0 * cm, 4.2 * cm, 8.0 * cm]))

    doc.build(story)
    out.seek(0)
    return out.getvalue()


def render_sidebar(profile: Dict[str, Any]) -> None:
    st.sidebar.header("Setup and privacy")
    st.sidebar.info(
        "This app is local/session-only. Nothing is stored after the app closes. "
        "To keep your data, download your JSON portfolio file and keep it safe."
    )

    with st.sidebar.expander("Registration / pathway questions", expanded=True):
        profile["psychologist_name"] = st.text_input(
            "Psychologist name",
            value=profile.get("psychologist_name", ""),
            key="profile_name",
        )
        profile["signature_name"] = st.text_input(
            "Signature name for exports",
            value=profile.get("signature_name", ""),
            key="profile_sig",
        )
        profile["registration_cycle_year_end"] = st.number_input(
            "Registration cycle year end (for 1 Dec previous year to 30 Nov this year)",
            min_value=2020,
            max_value=2100,
            step=1,
            value=int(profile.get("registration_cycle_year_end", date.today().year)),
            key="profile_cycle",
        )

        profile["full_year_general_registration"] = st.checkbox(
            "Held general registration for the full cycle",
            value=profile.get("full_year_general_registration", True),
            key="profile_full_year",
        )
        if profile["full_year_general_registration"]:
            profile["months_general_registration"] = 12
            st.caption("Full-year registration uses the standard 30-hour target.")
        else:
            profile["months_general_registration"] = st.number_input(
                "Full months of general registration in the cycle",
                min_value=0,
                max_value=12,
                value=int(profile.get("months_general_registration", 12)),
                step=1,
                key="profile_months",
            )
            st.caption(
                "Pro-rata target uses 2.5 hours total per full month, including 50 minutes of peer consultation per month."
            )

        profile["has_board_exemption"] = st.checkbox(
            "There is a Board exemption or variation affecting this cycle",
            value=profile.get("has_board_exemption", False),
            key="profile_exemption",
        )
        if profile["has_board_exemption"]:
            profile["exemption_notes"] = st.text_area(
                "Exemption / variation notes",
                value=profile.get("exemption_notes", ""),
                key="profile_exemption_notes",
            )

        profile["is_part_time"] = st.checkbox(
            "Practising part-time",
            value=profile.get("is_part_time", False),
            key="profile_part_time",
        )
        profile["has_endorsement"] = st.checkbox(
            "Has area of practice endorsement(s)",
            value=profile.get("has_endorsement", False),
            key="profile_has_endorsement",
        )
        if profile["has_endorsement"]:
            selected = st.multiselect(
                "Select endorsement area(s)",
                options=ENDORSEMENT_OPTIONS,
                default=profile.get("endorsements", []),
                key="profile_endorsements",
            )
            if "Other / custom" in selected:
                custom = st.text_input("Custom endorsement label", value="", key="profile_custom_endorsement")
                selected = [x for x in selected if x != "Other / custom"]
                if custom.strip():
                    selected.append(custom.strip())
            profile["endorsements"] = selected
        else:
            profile["endorsements"] = []

        profile["is_registrar"] = st.checkbox(
            "Currently in registrar program",
            value=profile.get("is_registrar", False),
            key="profile_registrar",
        )
        if profile["is_registrar"]:
            profile["registrar_area"] = st.text_input(
                "Registrar area",
                value=profile.get("registrar_area", ""),
                key="profile_registrar_area",
            )

        profile["is_board_approved_supervisor"] = st.checkbox(
            "Board-approved supervisor",
            value=profile.get("is_board_approved_supervisor", False),
            key="profile_supervisor",
        )
        if profile["is_board_approved_supervisor"]:
            profile["supervisor_last_training_date"] = st.text_input(
                "Last supervisor training/refresher date",
                value=profile.get("supervisor_last_training_date", ""),
                key="profile_supervisor_date",
                help="The guideline notes Board-approved supervisors must complete training or refresher at least every five years.",
            )

        profile["practice_context"] = st.text_area(
            "Practice context",
            value=profile.get("practice_context", ""),
            key="profile_context",
            help="Optional summary of your role, practice setting, or position.",
        )
        profile["position_description_note"] = st.text_area(
            "Position / practice relevance note",
            value=profile.get("position_description_note", ""),
            key="profile_position",
            help="Useful for showing how your learning plan relates to your current practice.",
        )


def render_dashboard(portfolio: Dict[str, Any]) -> None:
    metrics = compute_metrics(portfolio)
    st.subheader("Dashboard")
    st.caption(f"Tracking cycle: {metrics['cycle_label']}")

    c1, c2, c3 = st.columns(3)
    c1.metric("Total hours", f"{metrics['total_hours']} / {metrics['total_target']}", f"{metrics['total_remaining']} remaining")
    c2.metric("General CPD hours", f"{metrics['cpd_hours']} / {metrics['general_target']}", f"{metrics['cpd_remaining']} remaining")
    c3.metric("Peer consultation hours", f"{metrics['peer_hours']} / {metrics['peer_target']}", f"{metrics['peer_remaining']} remaining")

    st.info("Use decimal hours throughout the app. Example: 1.5 means 1 hour 30 minutes.")

    if portfolio["profile"].get("endorsements"):
        st.markdown("**Endorsement tracking**")
        rows = []
        targets = metrics["endorsement_targets"]
        for area, target in targets.items():
            actual = round(metrics["endorsement_hours"].get(area, 0.0), 2)
            rows.append({
                "Area": area,
                "Logged hours": actual,
                "Target / expectation": target,
                "Remaining": round(max(0.0, target - actual), 2),
            })
        st.dataframe(pd.DataFrame(rows), use_container_width=True, hide_index=True)

    warnings = []
    if metrics["peer_hours"] < metrics["peer_target"]:
        warnings.append("Peer consultation target is not yet met.")
    if metrics["cpd_hours"] < metrics["general_target"]:
        warnings.append("General CPD target is not yet met.")

    out_of_cycle_cpd = [
        x for x in portfolio["cpd_entries"]
        if x.get("date") and not in_cpd_cycle(x.get("date"), portfolio["profile"]["registration_cycle_year_end"])
    ]
    out_of_cycle_peer = [
        x for x in portfolio["peer_entries"]
        if x.get("date") and not in_cpd_cycle(x.get("date"), portfolio["profile"]["registration_cycle_year_end"])
    ]
    if out_of_cycle_cpd or out_of_cycle_peer:
        warnings.append("Some entries fall outside the selected CPD cycle and are not counted in the dashboard totals.")

    if warnings:
        for warning in warnings:
            st.warning(warning)
    else:
        st.success("Current logged hours meet the main Board minimum targets for this cycle.")

    generated = build_insights(portfolio, metrics)
    portfolio["summary_insights"] = st.text_area(
        "Summary insights",
        value=portfolio.get("summary_insights") or generated,
        height=180,
        help="Editable summary for your export documents.",
    )


def render_learning_plan(portfolio: Dict[str, Any]) -> None:
    st.subheader("Learning plan")
    st.caption("At least three goals are preloaded, but you can add as many as you need.")
    goals = portfolio["learning_goals"]

    if st.button("Add learning goal"):
        goals.append(
            {
                "id": new_id(),
                "title": f"Goal {len(goals) + 1}",
                "learning_need": "",
                "proposed_activities": "",
                "proposed_dates": "",
                "anticipated_outcomes": "",
                "review_date": "",
                "outcomes_achieved": "",
                "status": "Planned",
            }
        )

    delete_idx = None
    for idx, goal in enumerate(goals):
        with st.expander(goal.get("title") or f"Goal {idx + 1}", expanded=(idx < 3)):
            col1, col2 = st.columns([4, 1])
            goal["title"] = col1.text_input("Goal title", value=goal.get("title", ""), key=f"goal_title_{goal['id']}")
            goal["status"] = col2.selectbox(
                "Status",
                ["Planned", "In progress", "Reviewed", "Completed"],
                index=["Planned", "In progress", "Reviewed", "Completed"].index(goal.get("status", "Planned")),
                key=f"goal_status_{goal['id']}",
            )
            goal["learning_need"] = st.text_area(
                "Learning needs identified and goals set",
                value=goal.get("learning_need", ""),
                key=f"goal_need_{goal['id']}",
            )
            goal["proposed_activities"] = st.text_area(
                "Type of activities proposed to meet this need",
                value=goal.get("proposed_activities", ""),
                key=f"goal_activities_{goal['id']}",
            )
            goal["proposed_dates"] = st.text_input(
                "Dates proposed activities planned",
                value=goal.get("proposed_dates", ""),
                key=f"goal_dates_{goal['id']}",
            )
            goal["anticipated_outcomes"] = st.text_area(
                "Outcomes anticipated",
                value=goal.get("anticipated_outcomes", ""),
                key=f"goal_outcomes_{goal['id']}",
            )
            goal["review_date"] = st.text_input(
                "Review date",
                value=goal.get("review_date", ""),
                key=f"goal_review_{goal['id']}",
            )
            goal["outcomes_achieved"] = st.text_area(
                "Outcomes achieved",
                value=goal.get("outcomes_achieved", ""),
                key=f"goal_achieved_{goal['id']}",
            )
            if len(goals) > 3 and st.button("Delete this goal", key=f"delete_goal_{goal['id']}"):
                delete_idx = idx
    if delete_idx is not None:
        del goals[delete_idx]
        st.rerun()

    if goals:
        overview = pd.DataFrame(
            [
                {
                    "Goal": g.get("title", ""),
                    "Status": g.get("status", ""),
                    "Review date": g.get("review_date", ""),
                    "Planned activities": g.get("proposed_activities", ""),
                }
                for g in goals
            ]
        )
        st.dataframe(overview, use_container_width=True, hide_index=True)


def render_cpd_log(portfolio: Dict[str, Any]) -> None:
    st.subheader("General CPD log")
    goals = portfolio["learning_goals"]
    goal_options = goal_title_map(goals)
    entries = portfolio["cpd_entries"]

    if entries:
        rows = []
        for entry in entries:
            rows.append(
                {
                    "Date": entry.get("date", ""),
                    "Type": entry.get("activity_type", ""),
                    "Activity details": entry.get("activity_details", ""),
                    "Area": entry.get("area_of_practice", ""),
                    "Hours": entry.get("hours", 0.0),
                    "Goals": format_goal_links(entry.get("goal_ids", []), goal_options),
                    "Evidence": entry.get("evidence_type", ""),
                    "Reflection": entry.get("reflection", ""),
                }
            )
        st.dataframe(pd.DataFrame(rows), use_container_width=True, hide_index=True)

    st.markdown("**Add or edit CPD entry**")

    entry_options = {"New entry": None}
    for e in entries:
        label = f"{e.get('date', '')} | {e.get('activity_type', '')} | {safe_float(e.get('hours'))}h | {e.get('activity_details', '')[:40]}"
        entry_options[label] = e.get("id")

    cpd_labels = list(entry_options.keys())
    current_cpd_id = st.session_state.get("editing_cpd_id")
    current_cpd_label = "New entry"
    for label, value in entry_options.items():
        if value == current_cpd_id:
            current_cpd_label = label
            break
    current_cpd_index = cpd_labels.index(current_cpd_label) if current_cpd_label in cpd_labels else 0

    selected_label = st.selectbox(
        "Choose an existing CPD entry to edit, or leave on New entry",
        options=cpd_labels,
        index=current_cpd_index,
        key="cpd_select_entry",
    )

    selected_id = entry_options[selected_label]
    st.session_state.editing_cpd_id = selected_id
    selected_entry = get_entry_by_id(entries, selected_id)

    default_date = parse_iso_date(selected_entry.get("date")) if selected_entry else date.today()
    default_activity_type = selected_entry.get("activity_type", CPD_ACTIVITY_TYPES[0]) if selected_entry else CPD_ACTIVITY_TYPES[0]
    if default_activity_type not in CPD_ACTIVITY_TYPES:
        default_activity_type = "Other"
    default_hours = safe_float(selected_entry.get("hours")) if selected_entry else 0.0
    default_activity_details = selected_entry.get("activity_details", "") if selected_entry else ""
    default_area = selected_entry.get("area_of_practice", "") if selected_entry else ""
    default_goal_ids = selected_entry.get("goal_ids", []) if selected_entry else []
    default_reflection = selected_entry.get("reflection", "") if selected_entry else ""
    default_evidence_type = selected_entry.get("evidence_type", EVIDENCE_OPTIONS[0]) if selected_entry else EVIDENCE_OPTIONS[0]
    if default_evidence_type not in EVIDENCE_OPTIONS:
        default_evidence_type = "Other"
    default_evidence_details = selected_entry.get("evidence_details", "") if selected_entry else ""

    with st.form("cpd_form", clear_on_submit=False):
        c1, c2, c3 = st.columns(3)
        entry_date = c1.date_input("Date of activity", value=default_date)
        activity_type = c2.selectbox(
            "Type of activity",
            CPD_ACTIVITY_TYPES,
            index=CPD_ACTIVITY_TYPES.index(default_activity_type),
        )
        hours = c3.number_input(
            "Duration in decimal hours",
            min_value=0.0,
            step=0.25,
            format="%.2f",
            value=float(default_hours),
        )

        activity_details = st.text_area(
            "Activity details",
            value=default_activity_details,
            help="Name of course, presenter, institution, article, workshop details, etc.",
        )
        area = st.text_input("Area of practice (if applicable)", value=default_area)
        related_goals = st.multiselect(
            "Related learning goals",
            options=list(goal_options.keys()),
            default=default_goal_ids,
            format_func=lambda gid: goal_options[gid],
        )
        reflection = st.text_area(
            "Reflection",
            value=default_reflection,
            help="Written reflection is required. Keep this linked to what you learned and how it relates to practice and your learning plan.",
        )
        e1, e2 = st.columns([2, 3])
        evidence_type = e1.selectbox(
            "Evidence type",
            EVIDENCE_OPTIONS,
            index=EVIDENCE_OPTIONS.index(default_evidence_type),
        )
        evidence_details = e2.text_input(
            "Evidence detail / where stored",
            value=default_evidence_details,
            help="Describe the evidence and where you keep it. The app does not store attachments.",
        )

        col1, col2 = st.columns(2)
        save_clicked = col1.form_submit_button("Save CPD entry")
        delete_clicked = col2.form_submit_button("Delete selected CPD entry")

    if save_clicked:
        entry_id = selected_entry.get("id") if selected_entry else new_id()
        entry = {
            "id": entry_id,
            "date": entry_date.isoformat(),
            "activity_type": activity_type,
            "activity_details": activity_details,
            "area_of_practice": area,
            "hours": round(float(hours), 2),
            "goal_ids": related_goals,
            "reflection": reflection,
            "evidence_type": evidence_type,
            "evidence_details": evidence_details,
        }
        upsert_entry(entries, entry)
        st.session_state.editing_cpd_id = None
        st.success("CPD entry saved.")
        st.rerun()

    if delete_clicked:
        if selected_id:
            delete_entry(entries, selected_id)
            st.session_state.editing_cpd_id = None
            st.success("CPD entry deleted.")
            st.rerun()
        else:
            st.warning("Select an existing CPD entry first.")

def render_peer_log(portfolio: Dict[str, Any]) -> None:
    st.subheader("Peer consultation log")
    st.caption(
        "Enter the total duration and the own-practice component separately. The dashboard counts the own-practice hours toward the peer consultation target."
    )

    goals = portfolio["learning_goals"]
    goal_options = goal_title_map(goals)
    entries = portfolio["peer_entries"]

    if entries:
        rows = []
        for entry in entries:
            rows.append(
                {
                    "Date": entry.get("date", ""),
                    "Format": entry.get("format", ""),
                    "Focus": entry.get("focus", ""),
                    "Colleagues": entry.get("colleagues", ""),
                    "Total hours": entry.get("total_hours", 0.0),
                    "Own-practice hours": entry.get("own_practice_hours", 0.0),
                    "Area": entry.get("area_of_practice", ""),
                    "Goals": format_goal_links(entry.get("goal_ids", []), goal_options),
                    "Evidence": entry.get("evidence_type", ""),
                    "Reflection": entry.get("reflection", ""),
                }
            )
        st.dataframe(pd.DataFrame(rows), use_container_width=True, hide_index=True)

    st.markdown("**Add or edit peer consultation entry**")

    entry_options = {"New entry": None}
    for e in entries:
        label = f"{e.get('date', '')} | {e.get('format', '')} | {safe_float(e.get('own_practice_hours'))}h own-practice | {e.get('focus', '')[:40]}"
        entry_options[label] = e.get("id")

    peer_labels = list(entry_options.keys())
    current_peer_id = st.session_state.get("editing_peer_id")
    current_peer_label = "New entry"
    for label, value in entry_options.items():
        if value == current_peer_id:
            current_peer_label = label
            break
    current_peer_index = peer_labels.index(current_peer_label) if current_peer_label in peer_labels else 0

    selected_label = st.selectbox(
        "Choose an existing peer consultation entry to edit, or leave on New entry",
        options=peer_labels,
        index=current_peer_index,
        key="peer_select_entry",
    )

    selected_id = entry_options[selected_label]
    st.session_state.editing_peer_id = selected_id
    selected_entry = get_entry_by_id(entries, selected_id)

    default_date = parse_iso_date(selected_entry.get("date")) if selected_entry else date.today()
    default_format = selected_entry.get("format", PEER_FORMATS[0]) if selected_entry else PEER_FORMATS[0]
    if default_format not in PEER_FORMATS:
        default_format = "Other"
    default_focus = selected_entry.get("focus", "") if selected_entry else ""
    default_colleagues = selected_entry.get("colleagues", "") if selected_entry else ""
    default_total_hours = safe_float(selected_entry.get("total_hours")) if selected_entry else 0.0
    default_own_practice_hours = safe_float(selected_entry.get("own_practice_hours")) if selected_entry else 0.0
    default_area = selected_entry.get("area_of_practice", "") if selected_entry else ""
    default_goal_ids = selected_entry.get("goal_ids", []) if selected_entry else []
    default_reflection = selected_entry.get("reflection", "") if selected_entry else ""
    default_evidence_type = selected_entry.get("evidence_type", EVIDENCE_OPTIONS[0]) if selected_entry else EVIDENCE_OPTIONS[0]
    if default_evidence_type not in EVIDENCE_OPTIONS:
        default_evidence_type = "Other"
    default_evidence_details = selected_entry.get("evidence_details", "") if selected_entry else ""

    with st.form("peer_form", clear_on_submit=False):
        c1, c2 = st.columns(2)
        entry_date = c1.date_input("Date", value=default_date, key="peer_date")
        peer_format = c2.selectbox(
            "Peer consultation format",
            PEER_FORMATS,
            index=PEER_FORMATS.index(default_format),
        )
        focus = st.text_area(
            "Focus of peer consultation",
            value=default_focus,
            help="Topic, issue, problem, case focus, ethics, record keeping, management, burnout prevention, etc.",
        )
        colleagues = st.text_input("Colleagues involved", value=default_colleagues)
        c3, c4, c5 = st.columns(3)
        total_hours = c3.number_input(
            "Total duration in decimal hours",
            min_value=0.0,
            step=0.25,
            format="%.2f",
            value=float(default_total_hours),
        )
        own_practice_hours = c4.number_input(
            "Own-practice hours in decimal hours",
            min_value=0.0,
            step=0.25,
            format="%.2f",
            value=float(default_own_practice_hours),
        )
        area = c5.text_input("Area of practice", value=default_area)
        related_goals = st.multiselect(
            "Related learning goals",
            options=list(goal_options.keys()),
            default=default_goal_ids,
            format_func=lambda gid: goal_options[gid],
        )
        reflection = st.text_area(
            "Reflection",
            value=default_reflection,
            help="Written reflection is required for each peer consultation entry.",
        )
        e1, e2 = st.columns([2, 3])
        evidence_type = e1.selectbox(
            "Evidence type",
            EVIDENCE_OPTIONS,
            index=EVIDENCE_OPTIONS.index(default_evidence_type),
        )
        evidence_details = e2.text_input(
            "Evidence detail / where stored",
            value=default_evidence_details,
            help="Describe what evidence exists and where you keep it.",
        )

        col1, col2 = st.columns(2)
        save_clicked = col1.form_submit_button("Save peer consultation entry")
        delete_clicked = col2.form_submit_button("Delete selected peer entry")

    if save_clicked:
        if own_practice_hours > total_hours:
            st.error("Own-practice hours cannot exceed the total duration.")
        else:
            entry_id = selected_entry.get("id") if selected_entry else new_id()
            entry = {
                "id": entry_id,
                "date": entry_date.isoformat(),
                "format": peer_format,
                "focus": focus,
                "colleagues": colleagues,
                "total_hours": round(float(total_hours), 2),
                "own_practice_hours": round(float(own_practice_hours), 2),
                "area_of_practice": area,
                "goal_ids": related_goals,
                "reflection": reflection,
                "evidence_type": evidence_type,
                "evidence_details": evidence_details,
            }
            upsert_entry(entries, entry)
            st.session_state.editing_peer_id = None
            st.success("Peer consultation entry saved.")
            st.rerun()

    if delete_clicked:
        if selected_id:
            delete_entry(entries, selected_id)
            st.session_state.editing_peer_id = None
            st.success("Peer consultation entry deleted.")
            st.rerun()
        else:
            st.warning("Select an existing peer consultation entry first.")


def render_export(portfolio: Dict[str, Any]) -> None:
    st.subheader("Audit-style export")
    st.write(
        "Exports are structured around the guideline's Appendix E forms: learning plan, CPD activity log, peer consultation log, professional development journal, and peer consultation journal."
    )

    metrics = compute_metrics(portfolio)
    preview_text = portfolio.get("summary_insights") or build_insights(portfolio, metrics)
    st.text_area("Export preview summary", value=preview_text, height=160, disabled=True)

    docx_bytes = export_docx(portfolio)
    pdf_bytes = export_pdf(portfolio)

    base_name = portfolio["profile"].get("psychologist_name", "").strip().replace(" ", "_") or "psychology_cpd_portfolio"
    st.download_button(
        "Download Word export (.docx)",
        data=docx_bytes,
        file_name=f"{base_name}_audit_export.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    st.download_button(
        "Download PDF export (.pdf)",
        data=pdf_bytes,
        file_name=f"{base_name}_audit_export.pdf",
        mime="application/pdf",
    )

    st.info(
        "Store your external evidence files separately and keep them together with your Word/PDF export and your JSON portfolio file."
    )


def render_save_load(portfolio: Dict[str, Any]) -> None:
    st.subheader("Save / load portfolio")
    st.warning(
        "Confidentiality note: this app does not keep your data after the app closes. "
        "Your JSON portfolio file is your record. Keep it safe."
    )

    json_bytes = portfolio_json_bytes(portfolio)
    base_name = portfolio["profile"].get("psychologist_name", "").strip().replace(" ", "_") or "psychology_cpd_portfolio"
    st.download_button(
        "Download portfolio JSON",
        data=json_bytes,
        file_name=f"{base_name}.json",
        mime="application/json",
    )

    uploaded = st.file_uploader("Upload existing portfolio JSON", type=["json"])
    if uploaded is not None:
        try:
            loaded = load_portfolio_from_bytes(uploaded.read())
            if st.button("Load uploaded portfolio into this session"):
                st.session_state.portfolio = loaded
                st.session_state.editing_cpd_id = None
                st.session_state.editing_peer_id = None
                st.success("Portfolio loaded into the current session.")
                st.rerun()
        except Exception as e:
            st.error(f"Unable to load file: {e}")

    if st.button("Start a new blank portfolio"):
        st.session_state.portfolio = default_portfolio()
        st.session_state.editing_cpd_id = None
        st.session_state.editing_peer_id = None
        st.rerun()


def main() -> None:
    st.set_page_config(page_title=APP_TITLE, layout="wide")
    ensure_state()
    portfolio = st.session_state.portfolio

    st.title(APP_TITLE)
    st.caption(
        "Local-only CPD tracker for psychologists. Session-only by design: no automatic storage outside the JSON file you choose to save."
    )

    st.info(
        "This tool is designed around the uploaded guideline: 30 hours per cycle for full-year general registration, "
        "including at least 10 hours of peer consultation, plus a written learning plan, logs, reflections, and evidence records. "
        "The guideline allows electronic record-keeping and Appendix E provides form structures for audit-style documentation."
    )

    render_sidebar(portfolio["profile"])

    tabs = st.tabs(["Dashboard", "Learning plan", "CPD log", "Peer consultation", "Export", "Save / load"])

    with tabs[0]:
        render_dashboard(portfolio)
    with tabs[1]:
        render_learning_plan(portfolio)
    with tabs[2]:
        render_cpd_log(portfolio)
    with tabs[3]:
        render_peer_log(portfolio)
    with tabs[4]:
        render_export(portfolio)
    with tabs[5]:
        render_save_load(portfolio)

    st.divider()
    st.caption(
        "Privacy design: data is only kept in memory during the current session unless you explicitly download your JSON portfolio file."
    )


if __name__ == "__main__":
    main()